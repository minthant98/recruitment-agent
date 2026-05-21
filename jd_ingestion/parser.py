"""
parser.py — Extract raw text from a .docx or .pdf JD file.
"""
from pathlib import Path


def extract_text_from_docx(file_path: str) -> str:
    """
    Extract all text from a Word document (.docx) or PDF (.pdf).
    Returns a single string with newlines between paragraphs.
    Accepts either file type despite the function name (kept for compatibility).
    """
    path = Path(file_path)

    if not path.exists():
        raise FileNotFoundError(f"JD file not found: {file_path}")

    suffix = path.suffix.lower()

    if suffix == ".docx":
        return _extract_docx(path)
    elif suffix == ".pdf":
        return _extract_pdf(path)
    else:
        raise ValueError(f"Unsupported file type: {suffix}. Upload a .docx or .pdf file.")


def _extract_docx(path: Path) -> str:
    from docx import Document
    doc = Document(str(path))
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    return "\n".join(paragraphs)


def _extract_pdf(path: Path) -> str:
    try:
        import pypdf
        reader = pypdf.PdfReader(str(path))
        pages = []
        for page in reader.pages:
            text = page.extract_text()
            if text:
                pages.append(text.strip())
        return "\n".join(pages)
    except ImportError:
        # Fallback to pypdfium2 if pypdf not available
        import pypdfium2 as pdfium
        pdf    = pdfium.PdfDocument(str(path))
        pages  = []
        for page in pdf:
            textpage = page.get_textpage()
            text     = textpage.get_text_range()
            if text.strip():
                pages.append(text.strip())
        return "\n".join(pages)